"""
RAG ingestion pipeline — improved.

Key changes vs v1:
- Embeddings use a dedicated sentence-transformer model (nomic-embed-text via
  Ollama) instead of the chat LLM.  Embedding models are faster, cheaper, and
  produce better semantic vectors than generalist LLMs.
- Chunk size raised to 800 chars with 150 overlap — large enough to keep policy
  sentences together; overlap prevents answers being split across chunk boundaries.
- Separators extended to respect markdown headers and list items.
- Rich metadata added per chunk: source filename, chunk index, doc type, ingest
  timestamp.  Retriever can filter by these at query time.
- Deduplication: chunks whose content hash already exists in the store are
  skipped, so re-running ingest doesn't balloon the vector store.
- `ingest_documents` returns detailed per-file stats for the API response.
- `retrieve_documents` exposes MMR (Maximal Marginal Relevance) search to reduce
  redundant chunks in the returned context.
"""

from __future__ import annotations

import hashlib
import logging
import os
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import (
    DirectoryLoader,
    PyPDFLoader,
    TextLoader,
)
from langchain_community.vectorstores import Chroma
from langchain_ollama import OllamaEmbeddings

from app.config import settings

try:
    import docx as _docx
except ImportError:
    _docx = None

try:
    import openpyxl as _openpyxl
except ImportError:
    _openpyxl = None

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

_EMBED_MODEL = getattr(settings, "embedding_model", settings.llm_model)

_CHUNK_SIZE = 800
_CHUNK_OVERLAP = 150

_SEPARATORS = [
    "\n## ", "\n### ", "\n#### ",
    "\n\n",
    "\n- ", "\n* ",
    "\n",
    ". ", "! ", "? ",
    " ", "",
]


class DocumentIngestionPipeline:
    """Load, chunk, embed, and store documents for RAG retrieval."""

    def __init__(self):
        self.embeddings = OllamaEmbeddings(model=_EMBED_MODEL)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=_CHUNK_SIZE,
            chunk_overlap=_CHUNK_OVERLAP,
            separators=_SEPARATORS,
            length_function=len,
        )
        self.vector_store: Optional[Chroma] = None
        self._initialize_vector_store()

    # ------------------------------------------------------------------
    # Initialization
    # ------------------------------------------------------------------

    def _initialize_vector_store(self) -> None:
        persist_dir = settings.vector_store_path
        os.makedirs(persist_dir, exist_ok=True)
        self.vector_store = Chroma(
            persist_directory=persist_dir,
            embedding_function=self.embeddings,
        )
        if self._is_vector_store_empty():
            logger.info("Vector store empty — ingesting documents from data directory.")
            self.ingest_documents()

    def _is_vector_store_empty(self) -> bool:
        try:
            return len(self.vector_store.get(limit=1).get("ids", [])) == 0
        except Exception as exc:
            logger.warning("Could not inspect vector store: %s", exc)
            return True

    # ------------------------------------------------------------------
    # Loaders
    # ------------------------------------------------------------------

    def _load_docx(self, file_path: Path) -> List[Document]:
        if _docx is None:
            raise ImportError("python-docx is required for .docx files.")
        doc = _docx.Document(str(file_path))
        content = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if not content:
            return []
        return [Document(page_content=content, metadata={"source": str(file_path)})]

    def _load_xlsx(self, file_path: Path) -> List[Document]:
        if _openpyxl is None:
            raise ImportError("openpyxl is required for .xlsx files.")
        wb = _openpyxl.load_workbook(str(file_path), data_only=True)
        texts: List[str] = []
        for sheet in wb.worksheets:
            rows = [
                " ".join(str(c) for c in row if c is not None)
                for row in sheet.iter_rows(values_only=True)
            ]
            rows = [r.strip() for r in rows if r.strip()]
            if rows:
                texts.append(f"Sheet: {sheet.title}\n" + "\n".join(rows))
        content = "\n\n".join(texts).strip()
        if not content:
            return []
        return [Document(page_content=content, metadata={"source": str(file_path)})]

    # ------------------------------------------------------------------
    # Public: load
    # ------------------------------------------------------------------

    def load_documents(self, document_path: str = None) -> List[Document]:
        document_path = document_path or settings.document_path
        path = Path(document_path)
        if not path.exists():
            logger.warning("Document path does not exist: %s", document_path)
            return []

        docs: List[Document] = []

        for glob, loader_cls, label in [
            ("**/*.txt", TextLoader, "TXT"),
            ("**/*.pdf", PyPDFLoader, "PDF"),
        ]:
            try:
                loader = DirectoryLoader(document_path, glob=glob, loader_cls=loader_cls)
                loaded = loader.load()
                docs.extend(loaded)
                logger.info("Loaded %d %s file(s).", len(loaded), label)
            except Exception as exc:
                logger.error("Error loading %s files: %s", label, exc)

        for docx_path in path.rglob("*.docx"):
            try:
                docs.extend(self._load_docx(docx_path))
            except Exception as exc:
                logger.error("Error loading DOCX %s: %s", docx_path, exc)

        for xlsx_path in path.rglob("*.xlsx"):
            try:
                docs.extend(self._load_xlsx(xlsx_path))
            except Exception as exc:
                logger.error("Error loading XLSX %s: %s", xlsx_path, exc)

        return docs

    # ------------------------------------------------------------------
    # Public: ingest
    # ------------------------------------------------------------------

    def ingest_documents(self, document_path: str = None) -> Dict:
        documents = self.load_documents(document_path)
        if not documents:
            logger.warning("No documents found to ingest.")
            return {"status": "failed", "message": "No documents found."}

        chunks = self.text_splitter.split_documents(documents)
        logger.info("Split %d docs → %d chunks.", len(documents), len(chunks))

        now = datetime.now(timezone.utc).isoformat()
        existing_hashes = self._get_existing_hashes()
        new_chunks: List[Document] = []

        for idx, chunk in enumerate(chunks):
            content_hash = hashlib.md5(chunk.page_content.encode()).hexdigest()
            if content_hash in existing_hashes:
                continue

            source = chunk.metadata.get("source", "unknown")
            chunk.metadata.update({
                "chunk_index": idx,
                "content_hash": content_hash,
                "doc_type": Path(source).suffix.lstrip(".").lower(),
                "filename": Path(source).name,
                "ingested_at": now,
            })
            new_chunks.append(chunk)

        if not new_chunks:
            logger.info("All %d chunks already in store — nothing to add.", len(chunks))
            return {
                "status": "success",
                "documents_processed": len(documents),
                "chunks_created": 0,
                "chunks_skipped": len(chunks),
            }

        try:
            self.vector_store.add_documents(new_chunks)
            self.vector_store.persist()
            logger.info(
                "Ingested %d new chunks (%d skipped as duplicates).",
                len(new_chunks),
                len(chunks) - len(new_chunks),
            )
            return {
                "status": "success",
                "documents_processed": len(documents),
                "chunks_created": len(new_chunks),
                "chunks_skipped": len(chunks) - len(new_chunks),
            }
        except Exception as exc:
            logger.error("Error storing chunks: %s", exc)
            return {"status": "failed", "message": str(exc)}

    # ------------------------------------------------------------------
    # Public: retrieve
    # ------------------------------------------------------------------

    def retrieve_documents(
        self,
        query: str,
        k: int = 5,
        score_threshold: float = 0.30,
        use_mmr: bool = True,
        fetch_k: int = 20,
    ) -> List[Tuple[Document, float]]:
        if self.vector_store is None:
            logger.error("Vector store not initialised.")
            return []

        try:
            if use_mmr:

                try:

                    mmr_docs = self.vector_store.max_marginal_relevance_search(

                        query, k=k, fetch_k=fetch_k

                    )

                    scored = self.vector_store.similarity_search_with_score(query, k=fetch_k)

                    score_map = {doc.page_content: score for doc, score in scored}

                    results = [

                        (doc, score_map.get(doc.page_content, 0.0))

                        for doc in mmr_docs

                    ]

                except Exception as exc:

                    error_text = str(exc).lower()

                    if "not found" in error_text and _EMBED_MODEL != settings.llm_model:

                        logger.warning(

                            "Embedding model %s not found; retrying with LLM model %s.",

                            _EMBED_MODEL,

                            settings.llm_model,

                        )

                        self.embeddings = OllamaEmbeddings(model=settings.llm_model)

                        if hasattr(self.vector_store, "embedding_function"):

                            self.vector_store.embedding_function = self.embeddings

                        mmr_docs = self.vector_store.max_marginal_relevance_search(

                            query, k=k, fetch_k=fetch_k

                        )

                        scored = self.vector_store.similarity_search_with_score(query, k=fetch_k)

                        score_map = {doc.page_content: score for doc, score in scored}

                        results = [

                            (doc, score_map.get(doc.page_content, 0.0))

                            for doc in mmr_docs

                        ]

                    else:

                        raise

            else:

                try:
                    results = self.vector_store.similarity_search_with_score(query, k=k)
                except Exception as exc:
                    error_text = str(exc).lower()
                    if "not found" in error_text and _EMBED_MODEL != settings.llm_model:
                        logger.warning(
                            "Embedding model %s not found; retrying with LLM model %s.",
                            _EMBED_MODEL,
                            settings.llm_model,
                        )
                        self.embeddings = OllamaEmbeddings(model=settings.llm_model)
                        if hasattr(self.vector_store, "embedding_function"):
                            self.vector_store.embedding_function = self.embeddings
                        results = self.vector_store.similarity_search_with_score(query, k=k)
                    else:
                        raise

            filtered = [
                (doc, score) for doc, score in results if score <= score_threshold
            ]

            if not filtered:
                logger.info(
                    "All %d chunks exceeded score threshold %.2f — returning top %d anyway.",
                    len(results), score_threshold, min(k, len(results))
                )
                filtered = sorted(results, key=lambda x: x[1])[:k]

            logger.debug(
                "Retrieved %d chunks for query %r | scores: %s",
                len(filtered),
                query[:60],
                [round(s, 3) for _, s in filtered],
            )
            return filtered

        except Exception as exc:
            logger.error("Retrieval error: %s", exc)
            return []

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _get_existing_hashes(self) -> set:
        try:
            result = self.vector_store.get(include=["metadatas"])
            return {
                m.get("content_hash")
                for m in result.get("metadatas", [])
                if m and m.get("content_hash")
            }
        except Exception as exc:
            logger.warning("Could not fetch existing hashes: %s", exc)
            return set()


_rag_pipeline: Optional[DocumentIngestionPipeline] = None


def get_rag_pipeline() -> DocumentIngestionPipeline:
    global _rag_pipeline
    if _rag_pipeline is None:
        _rag_pipeline = DocumentIngestionPipeline()
    return _rag_pipeline
